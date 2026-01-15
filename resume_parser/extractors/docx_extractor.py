"""
DOCX text extraction module.

This module provides functionality to extract text content from
Microsoft Word (.docx) files using python-docx.
"""

from pathlib import Path
from typing import Union
import logging

from resume_parser.utils.text_utils import clean_text

logger = logging.getLogger(__name__)


class DOCXExtractor:
    """
    Extract text content from DOCX files.

    Uses python-docx for document parsing with support for
    paragraphs, tables, and headers.

    Example:
        extractor = DOCXExtractor()
        text = extractor.extract("resume.docx")
    """

    def __init__(self) -> None:
        """Initialize the DOCX extractor."""
        self._check_dependencies()

    def _check_dependencies(self) -> None:
        """Verify that python-docx is installed."""
        try:
            import docx  # noqa: F401
        except ImportError as e:
            raise ImportError(
                "python-docx is required for DOCX extraction. "
                "Install it with: pip install python-docx"
            ) from e

    def extract(self, file_path: Union[str, Path]) -> str:
        """
        Extract text from a DOCX file.

        Extracts text from paragraphs, tables, headers, and footers.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Extracted and cleaned text content.

        Raises:
            FileNotFoundError: If the file does not exist.
            ValueError: If the file is not a valid DOCX.
            RuntimeError: If text extraction fails.
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        if file_path.suffix.lower() != ".docx":
            raise ValueError(f"File is not a DOCX: {file_path}")

        return self._extract_text(file_path)

    def _extract_text(self, file_path: Path) -> str:
        """
        Perform the actual text extraction from DOCX.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Extracted text content.

        Raises:
            RuntimeError: If extraction fails.
        """
        import docx

        text_parts: list[str] = []

        try:
            document = docx.Document(file_path)

            # Extract headers
            for section in document.sections:
                header = section.header
                if header:
                    for paragraph in header.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text)

            # Extract main body paragraphs
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in document.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    text_parts.append(table_text)

            # Extract footers
            for section in document.sections:
                footer = section.footer
                if footer:
                    for paragraph in footer.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text)

        except Exception as e:
            raise RuntimeError(f"Failed to read DOCX file: {e}") from e

        if not text_parts:
            logger.warning(f"No text extracted from DOCX: {file_path}")
            return ""

        full_text = "\n".join(text_parts)
        return clean_text(full_text)

    def _extract_table_text(self, table) -> str:
        """
        Extract text from a DOCX table.

        Args:
            table: python-docx Table object.

        Returns:
            Extracted table text with proper formatting.
        """
        rows_text: list[str] = []

        for row in table.rows:
            cells_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cells_text.append(cell_text)
            if cells_text:
                rows_text.append(" | ".join(cells_text))

        return "\n".join(rows_text)

    def get_paragraph_count(self, file_path: Union[str, Path]) -> int:
        """
        Get the number of paragraphs in a DOCX file.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Number of non-empty paragraphs.

        Raises:
            FileNotFoundError: If the file does not exist.
            RuntimeError: If reading fails.
        """
        import docx

        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        try:
            document = docx.Document(file_path)
            return sum(1 for p in document.paragraphs if p.text.strip())
        except Exception as e:
            raise RuntimeError(f"Failed to read DOCX: {e}") from e

    def extract_metadata(self, file_path: Union[str, Path]) -> dict:
        """
        Extract metadata from a DOCX file.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Dictionary containing document metadata.

        Raises:
            FileNotFoundError: If the file does not exist.
            RuntimeError: If reading fails.
        """
        import docx

        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        try:
            document = docx.Document(file_path)
            props = document.core_properties

            return {
                "title": props.title,
                "author": props.author,
                "subject": props.subject,
                "keywords": props.keywords,
                "created": props.created,
                "modified": props.modified,
                "last_modified_by": props.last_modified_by,
                "category": props.category,
            }
        except Exception as e:
            raise RuntimeError(f"Failed to read DOCX metadata: {e}") from e

    def extract_with_formatting(self, file_path: Union[str, Path]) -> list[dict]:
        """
        Extract text with formatting information.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            List of dictionaries with text and formatting info.

        Raises:
            FileNotFoundError: If the file does not exist.
            RuntimeError: If reading fails.
        """
        import docx

        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        try:
            document = docx.Document(file_path)
            formatted_content: list[dict] = []

            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    formatted_content.append({
                        "text": paragraph.text,
                        "style": paragraph.style.name if paragraph.style else None,
                        "is_bold": any(run.bold for run in paragraph.runs),
                        "is_heading": (
                            paragraph.style
                            and "Heading" in paragraph.style.name
                        ),
                    })

            return formatted_content
        except Exception as e:
            raise RuntimeError(f"Failed to read DOCX: {e}") from e
